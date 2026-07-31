from docx import Document

def generate_word(requirements,solution,review):
    path="SolutionDocument.docx"
    doc=Document()
    doc.add_heading("Solution Document",0)
    doc.add_heading("Requirements",1)
    doc.add_paragraph(requirements)
    doc.add_heading("Solution",1)
    doc.add_paragraph(solution)
    doc.add_heading("Review",1)
    doc.add_paragraph(review)
    doc.save(path)
    return path
